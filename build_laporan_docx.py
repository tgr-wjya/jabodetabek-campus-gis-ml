import os

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 32, 67)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 73, 101)
    return p

def add_body_p(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(20, 20, 20)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p

def add_bullet_item(doc, bold_prefix, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_bullet = p.add_run("-  ")
    r_bullet.font.name = "Times New Roman"
    r_bullet.font.size = Pt(11)
    r_bullet.font.bold = True

    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(20, 20, 20)

    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p

def add_image_with_caption(doc, img_path, caption_text, width=None):
    if width is None:
        width = Inches(5.8)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(2)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=width)

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(8)
        run_cap = p_cap.add_run(caption_text)
        run_cap.font.name = "Times New Roman"
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        run_cap.font.color.rgb = RGBColor(90, 90, 90)

def main():
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 1. HALAMAN JUDUL
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    r_t1 = p_title.add_run("PROPOSAL INOVASI PEKAN INOVASI\n")
    r_t1.font.name = "Times New Roman"
    r_t1.font.size = Pt(16)
    r_t1.font.bold = True
    r_t1.font.color.rgb = RGBColor(15, 32, 67)

    r_t2 = p_title.add_run("GIS-EduTransit: Platform WebGIS Aksesibilitas Transportasi Massal dan Rekomendasi Lokasi Kampus Satelit Berbasis Machine Learning")
    r_t2.font.name = "Times New Roman"
    r_t2.font.size = Pt(13)
    r_t2.font.bold = True
    r_t2.font.color.rgb = RGBColor(27, 73, 101)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(18)
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("Proposal Karya Inovasi Pekan Inovasi Perguruan Tinggi\nUniversitas Mercu Buana")
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(70, 70, 70)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(2.5), Inches(4.0)]
    meta_data = [
        ("Nama Inovator / Tim", "TEGAR WIJAYA KUSUMA (NIM: 41523010217)"),
        ("Program Studi / Fakultas", "Teknik Informatika / Ilmu Komputer"),
        ("Afiliasi Perguruan Tinggi", "Universitas Mercu Buana, Jakarta"),
        ("Dosen Pembimbing", "MOHAMAD YUSUF, S.KOM., MCS"),
        ("Tahun Pelaksanaan", "2026")
    ]
    
    for i, (label, val) in enumerate(meta_data):
        row = table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = col_widths[0], col_widths[1]
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.name = "Times New Roman"
        r0.font.size = Pt(10.5)
        r0.font.bold = True
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(10.5)
        
        set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        
    set_table_borders(table, color="D0D0D0", sz="4")
    doc.add_page_break()

    # 2. RINGKASAN EKSEKUTIF
    add_heading_1(doc, "Ringkasan Eksekutif")
    add_body_p(doc, "Laporan ini menyajikan aplikasi WebGIS untuk dua keperluan utama: mengukur aksesibilitas transportasi umum pada 266 kampus di Jabodetabek dan Karawang, serta menentukan rekomendasi lokasi pembangunan kampus baru pada 299 kecamatan menggunakan algoritma Machine Learning Random Forest.")
    add_body_p(doc, "Aplikasi ini dapat dibuka langsung dari peramban web tanpa memerlukan instalasi perangkat lunak tambahan. Pemodelan Machine Learning mencapai akurasi 85%, di mana jumlah lulusan SMA dan jarak ke kawasan industri menjadi dua variabel utama dalam penentuan lokasi kampus baru.")

    # 3. LATAR BELAKANG
    add_heading_1(doc, "Latar Belakang")
    add_body_p(doc, "Persebaran kampus di wilayah Jabodetabek dan Karawang belum sepenuhnya seimbang dengan jalur transportasi umum seperti KRL, MRT, LRT, dan TransJakarta. Berdasarkan hasil pemetaan, 138 dari 266 kampus (51,9%) berada di luar jangkauan langsung stasiun atau halte utama.")
    add_body_p(doc, "Selain itu, penentuan lokasi kampus baru sering dilakukan tanpa pertimbangan data spasial yang jelas, seperti ketersediaan calon mahasiswa atau akses industri. Aplikasi ini dibuat untuk menyediakan data spasial yang jelas dan mudah diakses bagi pengambil keputusan.")

    # 4. TUJUAN INOVASI
    add_heading_1(doc, "Tujuan Inovasi")
    add_body_p(doc, "Tujuan dari pengembangan aplikasi ini adalah:")
    add_bullet_item(doc, "Peta Aksesibilitas: ", "Menampilkan posisi kampus terhadap jaringan stasiun dan halte angkutan umum secara interaktif.")
    add_bullet_item(doc, "Rekomendasi Lokasi: ", "Memberikan gambaran kecamatan yang berpotensi untuk lokasi kampus baru berbasis data.")
    add_bullet_item(doc, "Kemudahan Akses: ", "Menyediakan informasi spasial dalam format web yang mudah digunakan oleh masyarakat dan pengelola kampus.")

    # 5. DESKRIPSI INOVASI
    add_heading_1(doc, "Deskripsi Inovasi")
    add_body_p(doc, "Sistem ini terdiri dari dua modul:")
    add_bullet_item(doc, "Modul 1 (Aksesibilitas Kampus): ", "Membagi kampus menjadi dua kategori: Transit-Oriented (jarak <= 1.000m dari stasiun atau <= 500m dari halte) dan Transit-Isolated (di luar radius tersebut).")
    add_bullet_item(doc, "Modul 2 (Rekomendasi Lokasi): ", "Menggunakan Random Forest untuk mengelompokkan 299 kecamatan berdasarkan 5 variabel: jumlah lulusan SMA, jarak kawasan industri, akses jalan tol, luas wilayah, dan kepadatan kampus.")

    add_image_with_caption(doc, "screenshot/web/Default.png", "Gambar 1: Tampilan Utama WebGIS Aksesibilitas Kampus", width=Inches(5.8))
    add_image_with_caption(doc, "screenshot/visualizations/Peta_Zonasi_Rekomendasi_ML_Soal2.png", "Gambar 2: Peta Rekomendasi Lokasi Kampus Baru Berbasis Machine Learning", width=Inches(5.8))

    # 6. KEBARUAN DAN KEUNGGULAN INOVASI
    add_heading_1(doc, "Kebaruan dan Keunggulan Inovasi")
    add_bullet_item(doc, "Akses Web Ringan: ", "Dapat diakses langsung via browser HP atau laptop dengan pemuatan data yang cepat.")
    add_bullet_item(doc, "Berbasis Data Objektif: ", "Penentuan lokasi menggunakan data jumlah siswa dan kedekatan industri, bukan perkiraan subjektif.")
    add_bullet_item(doc, "Open Source: ", "Memanfaatkan teknologi open-source sehingga tidak memerlukan biaya lisensi perangkat lunak.")

    # 7. MANFAAT DAN DAMPAK
    add_heading_1(doc, "Manfaat dan Dampak")
    add_bullet_item(doc, "Calon Mahasiswa: ", "Memudahkan melihat kemudahan angkutan umum ke kampus yang dituju.")
    add_bullet_item(doc, "Pengelola Kampus & Pemerintah: ", "Menjadi rujukan untuk pembukaan rute angkutan pengumpan dan perencanaan lokasi kampus cabang baru.")

    # 8. TAHAPAN PENGEMBANGAN INOVASI
    add_heading_1(doc, "Tahapan Pengembangan Inovasi")
    add_bullet_item(doc, "1. Pengumpulan Data: ", "Mengumpulkan data titik kampus, stasiun, halte, dan batas kecamatan.")
    add_bullet_item(doc, "2. Pemrosesan Spasial: ", "Menghitung jarak dan kepadatan kampus per kecamatan di QGIS.")
    add_bullet_item(doc, "3. Pemodelan ML: ", "Melatih model Random Forest menggunakan data variabel kecamatan.")
    add_bullet_item(doc, "4. Peluncuran Web: ", "Membangun antarmuka WebGIS dan mengunggahnya ke server cloud.")

    # 9. POTENSI HILIRISASI DAN KEBERLANJUTAN
    add_heading_1(doc, "Potensi Hilirisasi dan Keberlanjutan")
    add_body_p(doc, "Metode pemetaan ini dapat direplikasi untuk wilayah perkotaan lainnya di Indonesia. Pengembangan selanjutnya dapat menambahkan data waktu tempuh lalu lintas secara langsung.")

    # 10. HAK KEKAYAAN INTELEKTUAL (HKI)
    add_heading_1(doc, "Hak Kekayaan Intelektual (HKI)")
    add_body_p(doc, "Informasi rencana perlindungan HKI untuk aplikasi ini:")
    add_bullet_item(doc, "Jenis Ciptaan: ", "Perangkat Lunak (WebGIS Dashboard)")
    add_bullet_item(doc, "Judul Ciptaan: ", "GIS-EduTransit & GeoRandomForest Dashboard")
    add_bullet_item(doc, "Pemegang Hak: ", "Tegar Wijaya Kusuma & Universitas Mercu Buana")
    add_bullet_item(doc, "Status: ", "Usulan Hak Cipta Perangkat Lunak pada DJKI Kemenkumham RI")

    # 11. TIM INOVATOR
    add_heading_1(doc, "Tim Inovator")
    add_bullet_item(doc, "Inovator Utama: ", "Tegar Wijaya Kusuma (NIM: 41523010217) - Mahasiswa Teknik Informatika Universitas Mercu Buana.")
    add_bullet_item(doc, "Dosen Pembimbing: ", "Mohamad Yusuf, S.Kom., MCS - Dosen Pengampu Mata Kuliah GIS.")

    # 12. BUKTI SUDAH TERDAFTARNYA CIPTAAN (SERTIFIKAT HAK KEKAYAAN INTELEKTUAL) PADA DJKI
    add_heading_1(doc, "Bukti sudah terdaftarnya Ciptaan (Sertifikat Hak Kekayaan Intelektual) pada DJKI")
    add_body_p(doc, "Dokumen usulan pendaftaran Hak Cipta Perangkat Lunak untuk platform WebGIS dan pemodelan rekomendasi lokasi pada Pangkalan Data Kekayaan Intelektual DJKI Kemenkumham RI.")

    add_image_with_caption(doc, "screenshot/Peta Aksesibilitas Transportasi Massal Perguruan Tinggi Jabodetabek.png", "Gambar 3: Lampiran Peta Cetak dan Dokumen Integrasi Ciptaan", width=Inches(5.8))

    # 13. PENUTUP
    add_heading_1(doc, "Penutup")
    add_body_p(doc, "Laporan inovasi ini menyajikan solusi praktis untuk memetakan aksesibilitas transportasi kampus dan memberikan rekomendasi lokasi kampus baru secara terukur. Diharapkan karya ini dapat berguna bagi perencanaan wilayah dan pendidikan tinggi.")

    output_path = "docs/Laporan_Pekan_Inovasi_GIS_E504.docx"
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    main()
